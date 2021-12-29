from random import sample
import re
from docx import Document
import time

def get15(chapter_sizes):
  chapters = list(range(0, len(chapter_sizes)))
  two_items_chapters = sample(chapters, 1)
  result = []
  for i in range(0, len(chapter_sizes)):
    if i in two_items_chapters:
      items = sample(list(range(0, chapter_sizes[i])), 10)
      items.sort()
      result.append(items)
    else:
      result.append(sample(list(range(0, chapter_sizes[i])), 1))
  return result


def load_exercises_21_22(file_name):
    exercise_tree = []
    current_chapter = []
    exercise_tree.append(current_chapter)
    current_exercise = None
    with open(file_name, 'r') as file:
        for line in file:
            if current_exercise is None:
                current_exercise = line
            else:
                if line != '\n':
                    current_exercise = current_exercise + line

    if current_exercise is not None:
        current_chapter.append(current_exercise)

    return exercise_tree


def load_exercises(file_name):
    exercise_tree = []
    current_chapter = None
    current_exercise = None
    with open(file_name, 'r') as file:
        for line in file:
            if re.search('^Capitolo [IVX.]+ Esercizi', line):
                if current_exercise is not None:
                    current_chapter.append(current_exercise)
                    current_exercise = None
                current_chapter = []
                exercise_tree.append(current_chapter)
            elif re.search('^[0-9]+', line):
                if current_exercise is not None:
                    current_chapter.append(current_exercise)
                current_exercise = line
            else:
                if line != '\n':
                    current_exercise = current_exercise + line

    if current_exercise is not None:
        current_chapter.append(current_exercise)

    return exercise_tree


def create_document(exercises, selection):
    document = Document()
    for i in range(0, len(selection)):
        add_scheda_to_document(document, exercises, i, selection[i])
    return document


def add_scheda_to_document(document, exercises, scheda, exercise):
    document.add_heading('Scheda n. ' + str(scheda + 1))
    for i in range(0, len(exercise)):
        add_exercise_to_document(document, exercises[i], i, exercise[i])
    document.add_page_break()


def add_exercise_to_document(document, chapter_texts, chapter_index, tests):
    for test in tests:
        document.add_heading('Esercizio n. ' + str(chapter_index + 1) + '.' + str(test + 1), level=2)
        document.add_paragraph(chapter_texts[test])
        document.add_paragraph('')


def create_schede(exercises, students_number):
    chapter_sizes = [len(chapter) for chapter in exercises]
    selection = []
    for i in range(students_number):
        selection.append(get15(chapter_sizes))
    document = create_document(exercises, selection)
    return document
